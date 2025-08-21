"""
Attachment Processor for Construction Document Analysis

Processes various file types from email attachments including PDFs, Excel spreadsheets,
Word documents, and images to extract construction-related data.
"""

import io
import logging
import mimetypes
from typing import Dict, List, Any, Optional, Union, Tuple
from dataclasses import dataclass
from enum import Enum
from pathlib import Path

try:
    import PyPDF2
    from openpyxl import load_workbook
    from docx import Document
    from PIL import Image, ImageEnhance
    import pytesseract
    PDF_PROCESSING_AVAILABLE = True
except ImportError:
    PDF_PROCESSING_AVAILABLE = False
    logging.warning("Document processing libraries not available. Attachment processing will be limited.")

logger = logging.getLogger(__name__)


class FileType(Enum):
    """Supported file types for processing."""
    PDF = "pdf"
    EXCEL = "excel"
    WORD = "word"
    IMAGE = "image"
    TEXT = "text"
    UNKNOWN = "unknown"


@dataclass
class AttachmentData:
    """Structured data extracted from attachments."""
    filename: str
    file_type: FileType
    content: str = ""
    extracted_data: Dict[str, Any] = None
    metadata: Dict[str, Any] = None
    processing_status: str = "success"
    error_message: str = ""
    
    def __post_init__(self):
        if self.extracted_data is None:
            self.extracted_data = {}
        if self.metadata is None:
            self.metadata = {}


class MockAttachmentProcessor:
    """Mock attachment processor for development when libraries are not available."""
    
    async def initialize(self):
        logger.warning("Using mock attachment processor - document libraries not available")
    
    async def process_attachment(self, filename: str, content: bytes, project_id: str = None) -> Optional[Dict[str, Any]]:
        file_type = self._detect_file_type(filename)
        
        return {
            "file_type": file_type.value,
            "data": {
                "content": f"Mock processed content for {filename}",
                "type": "mock_data",
                "construction_keywords": ["project", "schedule", "budget"],
                "confidence": 0.5
            },
            "metadata": {
                "filename": filename,
                "size": len(content),
                "processed_at": "2024-01-01T00:00:00"
            }
        }
    
    def _detect_file_type(self, filename: str) -> FileType:
        ext = Path(filename).suffix.lower()
        
        if ext == '.pdf':
            return FileType.PDF
        elif ext in ['.xlsx', '.xls']:
            return FileType.EXCEL
        elif ext in ['.docx', '.doc']:
            return FileType.WORD
        elif ext in ['.png', '.jpg', '.jpeg', '.gif', '.bmp']:
            return FileType.IMAGE
        elif ext in ['.txt', '.csv']:
            return FileType.TEXT
        else:
            return FileType.UNKNOWN


class AttachmentProcessor:
    """
    Processes email attachments to extract construction-related data.
    
    Supports various file formats and uses AI-powered analysis to identify
    construction project information, schedules, budgets, and other relevant data.
    """
    
    def __init__(self):
        self.supported_types = {
            '.pdf': FileType.PDF,
            '.xlsx': FileType.EXCEL,
            '.xls': FileType.EXCEL,
            '.docx': FileType.WORD,
            '.doc': FileType.WORD,
            '.png': FileType.IMAGE,
            '.jpg': FileType.IMAGE,
            '.jpeg': FileType.IMAGE,
            '.gif': FileType.IMAGE,
            '.bmp': FileType.IMAGE,
            '.tiff': FileType.IMAGE,
            '.txt': FileType.TEXT,
            '.csv': FileType.TEXT
        }
        
        # Construction-related keywords for content analysis
        self.construction_keywords = {
            'project': ['project', 'site', 'construction', 'building', 'development'],
            'schedule': ['schedule', 'timeline', 'deadline', 'milestone', 'phase', 'duration'],
            'budget': ['budget', 'cost', 'price', 'estimate', 'expense', 'payment'],
            'rfi': ['rfi', 'request for information', 'clarification', 'question'],
            'change_order': ['change order', 'co', 'modification', 'amendment'],
            'materials': ['material', 'supply', 'delivery', 'shipment', 'equipment'],
            'inspection': ['inspection', 'quality', 'check', 'review', 'audit'],
            'safety': ['safety', 'hazard', 'incident', 'accident', 'osha']
        }
    
    async def initialize(self):
        """Initialize the attachment processor."""
        if not PDF_PROCESSING_AVAILABLE:
            logger.warning("Document processing libraries not available")
        else:
            logger.info("AttachmentProcessor initialized successfully")
    
    async def process_attachment(self, filename: str, content: bytes, project_id: str = None) -> Optional[Dict[str, Any]]:
        """
        Process an email attachment and extract relevant data.
        
        Args:
            filename: Name of the attachment file
            content: Binary content of the file
            project_id: Optional project ID for context
            
        Returns:
            Dictionary with extracted data and metadata, or None if processing failed
        """
        if not PDF_PROCESSING_AVAILABLE:
            mock_processor = MockAttachmentProcessor()
            await mock_processor.initialize()
            return await mock_processor.process_attachment(filename, content, project_id)
        
        try:
            # Detect file type
            file_type = self._detect_file_type(filename)
            
            if file_type == FileType.UNKNOWN:
                logger.warning(f"Unsupported file type: {filename}")
                return None
            
            # Process based on file type
            attachment_data = AttachmentData(
                filename=filename,
                file_type=file_type,
                metadata={
                    "size": len(content),
                    "project_id": project_id,
                    "processed_at": "2024-01-01T00:00:00"  # Would use datetime.now().isoformat()
                }
            )
            
            if file_type == FileType.PDF:
                await self._process_pdf(content, attachment_data)
            elif file_type == FileType.EXCEL:
                await self._process_excel(content, attachment_data)
            elif file_type == FileType.WORD:
                await self._process_word(content, attachment_data)
            elif file_type == FileType.IMAGE:
                await self._process_image(content, attachment_data)
            elif file_type == FileType.TEXT:
                await self._process_text(content, attachment_data)
            
            # Analyze content for construction data
            if attachment_data.content:
                await self._analyze_construction_content(attachment_data)
            
            return {
                "file_type": attachment_data.file_type.value,
                "data": attachment_data.extracted_data,
                "metadata": attachment_data.metadata
            }
            
        except Exception as e:
            logger.error(f"Failed to process attachment {filename}: {e}")
            return {
                "file_type": "error",
                "data": {},
                "metadata": {"error": str(e), "filename": filename}
            }
    
    def _detect_file_type(self, filename: str) -> FileType:
        """Detect file type from filename extension."""
        ext = Path(filename).suffix.lower()
        return self.supported_types.get(ext, FileType.UNKNOWN)
    
    async def _process_pdf(self, content: bytes, attachment_data: AttachmentData):
        """Process PDF file and extract text content."""
        try:
            pdf_file = io.BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            # Extract text from all pages
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                text_content.append(page.extract_text())
            
            attachment_data.content = '\n'.join(text_content)
            
            # Extract PDF metadata
            if pdf_reader.metadata:
                attachment_data.metadata.update({
                    "pages": len(pdf_reader.pages),
                    "title": pdf_reader.metadata.get("/Title", ""),
                    "author": pdf_reader.metadata.get("/Author", ""),
                    "subject": pdf_reader.metadata.get("/Subject", "")
                })
            
            logger.info(f"Extracted text from PDF: {attachment_data.filename} ({len(pdf_reader.pages)} pages)")
            
        except Exception as e:
            logger.error(f"Failed to process PDF {attachment_data.filename}: {e}")
            attachment_data.processing_status = "error"
            attachment_data.error_message = str(e)
    
    async def _process_excel(self, content: bytes, attachment_data: AttachmentData):
        """Process Excel file and extract structured data."""
        try:
            excel_file = io.BytesIO(content)
            workbook = load_workbook(excel_file, data_only=True)
            
            extracted_data = {}
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_data = []
                
                # Extract data from each row
                for row in sheet.iter_rows(values_only=True):
                    if any(cell is not None for cell in row):
                        row_data = [str(cell) if cell is not None else "" for cell in row]
                        sheet_data.append(row_data)
                        text_content.append(" ".join(row_data))
                
                extracted_data[sheet_name] = sheet_data
            
            attachment_data.content = '\n'.join(text_content)
            attachment_data.extracted_data["excel_data"] = extracted_data
            attachment_data.metadata["sheets"] = list(workbook.sheetnames)
            
            logger.info(f"Extracted data from Excel: {attachment_data.filename} ({len(workbook.sheetnames)} sheets)")
            
        except Exception as e:
            logger.error(f"Failed to process Excel {attachment_data.filename}: {e}")
            attachment_data.processing_status = "error"
            attachment_data.error_message = str(e)
    
    async def _process_word(self, content: bytes, attachment_data: AttachmentData):
        """Process Word document and extract text content."""
        try:
            word_file = io.BytesIO(content)
            doc = Document(word_file)
            
            # Extract text from paragraphs
            text_content = []
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text.strip() for cell in row.cells]
                    table_data.append(row_data)
                    text_content.append(" ".join(row_data))
                tables_data.append(table_data)
            
            attachment_data.content = '\n'.join(text_content)
            if tables_data:
                attachment_data.extracted_data["tables"] = tables_data
            
            attachment_data.metadata.update({
                "paragraphs": len(doc.paragraphs),
                "tables": len(doc.tables)
            })
            
            logger.info(f"Extracted text from Word document: {attachment_data.filename}")
            
        except Exception as e:
            logger.error(f"Failed to process Word document {attachment_data.filename}: {e}")
            attachment_data.processing_status = "error"
            attachment_data.error_message = str(e)
    
    async def _process_image(self, content: bytes, attachment_data: AttachmentData):
        """Process image file using OCR to extract text."""
        try:
            image_file = io.BytesIO(content)
            image = Image.open(image_file)
            
            # Enhance image for better OCR
            enhancer = ImageEnhance.Contrast(image)
            enhanced_image = enhancer.enhance(2.0)
            
            # Perform OCR
            try:
                text = pytesseract.image_to_string(enhanced_image)
                attachment_data.content = text.strip()
            except Exception as ocr_error:
                logger.warning(f"OCR failed for {attachment_data.filename}: {ocr_error}")
                attachment_data.content = ""
            
            # Extract image metadata
            attachment_data.metadata.update({
                "width": image.width,
                "height": image.height,
                "format": image.format,
                "mode": image.mode
            })
            
            logger.info(f"Processed image: {attachment_data.filename} ({image.width}x{image.height})")
            
        except Exception as e:
            logger.error(f"Failed to process image {attachment_data.filename}: {e}")
            attachment_data.processing_status = "error"
            attachment_data.error_message = str(e)
    
    async def _process_text(self, content: bytes, attachment_data: AttachmentData):
        """Process plain text or CSV files."""
        try:
            # Try to decode as UTF-8, fallback to latin-1
            try:
                text_content = content.decode('utf-8')
            except UnicodeDecodeError:
                text_content = content.decode('latin-1', errors='ignore')
            
            attachment_data.content = text_content
            
            # If it's a CSV, try to parse structure
            if attachment_data.filename.lower().endswith('.csv'):
                lines = text_content.split('\n')
                csv_data = []
                for line in lines:
                    if line.strip():
                        row = [cell.strip().strip('"') for cell in line.split(',')]
                        csv_data.append(row)
                
                attachment_data.extracted_data["csv_data"] = csv_data
                attachment_data.metadata["rows"] = len(csv_data)
            
            logger.info(f"Processed text file: {attachment_data.filename}")
            
        except Exception as e:
            logger.error(f"Failed to process text file {attachment_data.filename}: {e}")
            attachment_data.processing_status = "error"
            attachment_data.error_message = str(e)
    
    async def _analyze_construction_content(self, attachment_data: AttachmentData):
        """Analyze content for construction-related information."""
        try:
            content_lower = attachment_data.content.lower()
            
            # Find construction keywords
            found_categories = {}
            for category, keywords in self.construction_keywords.items():
                matches = []
                for keyword in keywords:
                    if keyword in content_lower:
                        matches.append(keyword)
                
                if matches:
                    found_categories[category] = matches
            
            # Calculate relevance score
            total_keywords = sum(len(keywords) for keywords in self.construction_keywords.values())
            found_keywords = sum(len(matches) for matches in found_categories.values())
            relevance_score = found_keywords / total_keywords if total_keywords > 0 else 0.0
            
            # Extract potential project identifiers
            project_patterns = self._extract_project_patterns(attachment_data.content)
            
            # Extract dates and numbers
            dates = self._extract_dates(attachment_data.content)
            numbers = self._extract_numbers(attachment_data.content)
            
            # Store analysis results
            attachment_data.extracted_data.update({
                "construction_relevance": relevance_score,
                "construction_categories": found_categories,
                "project_patterns": project_patterns,
                "dates": dates,
                "numbers": numbers,
                "word_count": len(attachment_data.content.split())
            })
            
            logger.info(f"Construction analysis complete for {attachment_data.filename}: {relevance_score:.2f} relevance")
            
        except Exception as e:
            logger.error(f"Failed to analyze construction content: {e}")
    
    def _extract_project_patterns(self, content: str) -> List[str]:
        """Extract potential project identifiers from content."""
        import re
        
        patterns = []
        
        # Project number patterns (e.g., PRJ-2024-001, Project 12345)
        project_numbers = re.findall(r'(?:project|prj|job)\s*[#:-]?\s*(\w+[-]?\d+)', content, re.IGNORECASE)
        patterns.extend(project_numbers)
        
        # Site/building names
        site_patterns = re.findall(r'(?:site|building|facility)\s+(\w+(?:\s+\w+)*)', content, re.IGNORECASE)
        patterns.extend([match for match in site_patterns if len(match) > 3])
        
        return list(set(patterns))[:10]  # Limit to 10 patterns
    
    def _extract_dates(self, content: str) -> List[str]:
        """Extract dates from content."""
        import re
        
        # Common date patterns
        date_patterns = [
            r'\d{1,2}[/-]\d{1,2}[/-]\d{2,4}',  # MM/DD/YYYY or DD/MM/YYYY
            r'\d{4}[/-]\d{1,2}[/-]\d{1,2}',    # YYYY/MM/DD
            r'(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec)\s+\d{1,2},?\s+\d{4}',  # Month DD, YYYY
        ]
        
        dates = []
        for pattern in date_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            dates.extend(matches)
        
        return list(set(dates))[:20]  # Limit to 20 dates
    
    def _extract_numbers(self, content: str) -> Dict[str, List[str]]:
        """Extract currency and measurement numbers from content."""
        import re
        
        numbers = {
            "currency": [],
            "measurements": [],
            "percentages": []
        }
        
        # Currency patterns
        currency_patterns = [
            r'\$[\d,]+(?:\.\d{2})?',  # $1,234.56
            r'USD\s*[\d,]+(?:\.\d{2})?',  # USD 1234.56
        ]
        
        for pattern in currency_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            numbers["currency"].extend(matches)
        
        # Measurement patterns
        measurement_patterns = [
            r'\d+(?:\.\d+)?\s*(?:ft|feet|inches?|in|meters?|m|yards?|yd)',
            r'\d+(?:\.\d+)?\s*(?:sq\.?\s*ft|square\s+feet|sf)',
            r'\d+(?:\.\d+)?\s*(?:lbs?|pounds?|kg|kilograms?|tons?)'
        ]
        
        for pattern in measurement_patterns:
            matches = re.findall(pattern, content, re.IGNORECASE)
            numbers["measurements"].extend(matches)
        
        # Percentage patterns
        percentage_matches = re.findall(r'\d+(?:\.\d+)?%', content)
        numbers["percentages"].extend(percentage_matches)
        
        # Limit results
        for key in numbers:
            numbers[key] = list(set(numbers[key]))[:10]
        
        return numbers


# Export the appropriate class based on availability
if PDF_PROCESSING_AVAILABLE:
    AttachmentProcessorImpl = AttachmentProcessor
else:
    AttachmentProcessorImpl = MockAttachmentProcessor

# Alias for easier import
AttachmentProcessor = AttachmentProcessorImpl